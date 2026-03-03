from fastapi                 import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic                import BaseModel
from ultralytics             import YOLO
from datetime                import datetime
from typing                  import List,Dict
from docxtpl import DocxTemplate
from pathlib import Path
from docx2pdf import convert
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess, uuid, json, tempfile, base64, io
import uvicorn
import cv2
import numpy  as np
import base64, os, time, csv


MODEL_PATH  = "best.pt"
CLASS_NAMES = ["Normal", "Bulging", "Herniation"]
DISC_LEVELS = ["L1-L2", "L2-L3", "L3-L4", "L4-L5", "L5-S1"]
CONF_THRESH = 0.25

TEMPLATE_PATH = "report_template.docx"
TEMP_DIR      = Path(tempfile.gettempdir()) / "spine_reports"
TEMP_DIR.mkdir(exist_ok=True)

RESULTS_CSV = "results.csv" 

CLASS_CONFIG = {
    0: {"name": "Normal",     "color": (0,  200,   0), "severity": "low",      "emoji": "✅"},
    1: {"name": "Bulging",    "color": (0,  200, 255), "severity": "moderate", "emoji": "⚠️"},
    2: {"name": "Herniation", "color": (0,   0,  255), "severity": "severe",   "emoji": "🔴"},
}


CONDITION_STYLES = {
    "Normal"      : {"text": RGBColor(0,   150,   0),  "bg": "E6F7EE"},
    "Bulging"     : {"text": RGBColor(0,   150, 200),  "bg": "E6F7FF"},
    "Herniation"  : {"text": RGBColor(180,   0,   0),  "bg": "FDF0EF"},
    "Not Detected": {"text": RGBColor(100, 100, 100),  "bg": "F4F5F7"},
}

#  RESPONSE SCHEMAS

class MetricsResponse(BaseModel):
    map50      : float       
    map50_95   : float       
    precision  : float       
    recall     : float       
    per_class  : Dict[str, float]   
    source     : str    

class DiscResult(BaseModel):
    disc_level : str    
    condition  : str    
    confidence : float  
    severity   : str    

class Report(BaseModel):
    image_name      : str
    timestamp       : str
    discs           : List[DiscResult]
    summary         : dict
    overall_status  : str
    processing_time : float

class FullResponse(BaseModel):
    report          : Report
    annotated_image : str   


#  APP SETUP

app = FastAPI(
    title       = "Spinal Cord Damage Detection API",
    description = "Upload a lumbar spine MRI → JSON report + base64 annotated image",
    version     = "1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins     = ["*"],  
    allow_credentials = True,
    allow_methods     = ["*"],
    allow_headers     = ["*"],
)

model = None

@app.on_event("startup")
def load_model():
    global model
    if not os.path.exists(MODEL_PATH):
        print(f"⚠️  best.pt not found at '{MODEL_PATH}'")
        print("    Copy your trained weights file here and restart.")
        return
    model = YOLO(MODEL_PATH)
    print(f"✅ Model loaded → {MODEL_PATH}")



#  HELPER FUNCTIONS

def decode_image(file_bytes: bytes) -> np.ndarray:
    arr = np.frombuffer(file_bytes, np.uint8)
    img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    return img


def run_inference(img: np.ndarray) -> list:
    if model is None:
        raise HTTPException(
            status_code = 503,
            detail      = "Model not loaded. Copy best.pt here and restart."
        )

    # ── Model runs here — detection happens once 
    results = model.predict(
        source  = img,
        conf    = CONF_THRESH,
        imgsz   = 640,
        verbose = False,
    )
    result = results[0]

    # ── Read coordinates from model output 
    detections = []
    if result.boxes and len(result.boxes) > 0:
        for box in result.boxes:
            x1, y1, x2, y2 = map(int, box.xyxy[0])
            detections.append({
                "cls_id" : int(box.cls[0]),
                "conf"   : round(float(box.conf[0]), 4),
                "x1": x1, "y1": y1,
                "x2": x2, "y2": y2,
                "cy"     : (y1 + y2) // 2,  
            })

    detections = sorted(detections, key=lambda d: d["cy"])
    for i, det in enumerate(detections):
        det["disc_level"] = DISC_LEVELS[i] if i < len(DISC_LEVELS) else f"Disc-{i+1}"

    return detections


def draw_boxes(img: np.ndarray, detections: list) -> np.ndarray:
    vis  = img.copy()
    h, w = vis.shape[:2]
    font = cv2.FONT_HERSHEY_SIMPLEX

    # ── Bounding boxes + labels 
    for det in detections:
        cfg   = CLASS_CONFIG.get(det["cls_id"], {"name": "Unknown", "color": (180, 180, 180)})
        color = cfg["color"]
        label = f"{det['disc_level']}  |  {cfg['name']}  {det['conf']:.2f}"
        x1, y1, x2, y2 = det["x1"], det["y1"], det["x2"], det["y2"]

        # Box
        cv2.rectangle(vis, (x1, y1), (x2, y2), color, 2)

        # Label background
        (tw, th), _ = cv2.getTextSize(label, font, 0.55, 2)
        lbl_y1 = max(y1 - th - 10, 0)
        cv2.rectangle(vis, (x1, lbl_y1), (x1 + tw + 8, lbl_y1 + th + 8), color, -1)

        # Label text
        cv2.putText(vis, label, (x1 + 4, lbl_y1 + th + 2),
                    font, 0.55, (255, 255, 255), 2, cv2.LINE_AA)

    # ── Colour legend (bottom left) 
    n      = len(CLASS_CONFIG)
    box_h  = n * 24 + 30
    x0, y0 = 10, h - box_h - 10
    cv2.rectangle(vis, (x0-5, y0-22), (x0+220, y0+box_h-8), (20, 20, 20), -1)
    cv2.putText(vis, "LEGEND", (x0, y0-5), font, 0.5, (255, 255, 255), 1)
    for i, (cls_id, cfg) in enumerate(CLASS_CONFIG.items()):
        y = y0 + 18 + i * 24
        cv2.rectangle(vis, (x0, y-13), (x0+16, y+3), cfg["color"], -1)
        cv2.putText(vis, cfg["name"], (x0+22, y), font, 0.48, (220, 220, 220), 1)

    # ── Title bar (top) 
    cv2.rectangle(vis, (0, 0), (w, 30), (20, 20, 20), -1)
    cv2.putText(vis, "LUMBAR SPINE DAMAGE DETECTION", (10, 21),
                font, 0.65, (255, 255, 255), 2, cv2.LINE_AA)

    return vis  


def numpy_to_base64(vis: np.ndarray) -> str:
    _, buffer  = cv2.imencode(".jpg", vis)
    b64_bytes  = base64.b64encode(buffer.tobytes())
    b64_string = b64_bytes.decode("utf-8")
    return f"data:image/jpeg;base64,{b64_string}"


def build_report(detections: list, filename: str, elapsed: float) -> dict:
    detected_map = {d["disc_level"]: d for d in detections}
    summary      = {"Normal": 0, "Bulging": 0, "Herniation": 0, "Not_Detected": 0}
    disc_rows    = []

    for level in DISC_LEVELS:
        if level in detected_map:
            det  = detected_map[level]
            cfg  = CLASS_CONFIG.get(det["cls_id"], {"name": "Unknown", "severity": "unknown"})
            cond = cfg["name"]
            sev  = cfg["severity"]
            conf = det["conf"]
            summary[cond] = summary.get(cond, 0) + 1
        else:
            cond = "Not Detected"
            sev  = "unknown"
            conf = 0.0
            summary["Not_Detected"] += 1

        disc_rows.append({
            "disc_level" : level,
            "condition"  : cond,
            "confidence" : conf,
            "severity"   : sev,
        })

    # Overall status based on worst finding
    if summary.get("Herniation", 0) > 0:
        overall = "Critical"
    elif summary.get("Bulging", 0) > 0:
        overall = "Attention Required"
    else:
        overall = "Normal"

    return {
        "image_name"      : filename,
        "timestamp"       : datetime.utcnow().isoformat(),
        "discs"           : disc_rows,
        "summary"         : summary,
        "overall_status"  : overall,
        "processing_time" : round(elapsed, 3),
    }

def load_metrics_from_csv(csv_path: str) -> dict:

    # ── Check file exists
    if not os.path.exists(csv_path):
        raise HTTPException(
            status_code = 404,
            detail      = (
                f"results.csv not found at '{csv_path}'. "
                f"Place it in the same folder as metrics.py."
            )
        )

    # ── Read all lines from file 
    with open(csv_path, "r") as f:
        lines = [line.strip() for line in f.readlines()]

    # ── Remove empty lines 
    lines = [line for line in lines if line]

    # ── Parse both sections 
    overall_metrics = {}
    per_class       = {}
    current_section = None  

    for line in lines:
        parts = [p.strip() for p in line.split(",")]

        # ── Detect section headers 
        if parts[0] == "Metric" and parts[1] == "Value":
            current_section = "overall"
            continue

        if parts[0] == "Class Name" and parts[1] == "AP@50":
            current_section = "perclass"
            continue

        # ── Read overall metrics section 
        if current_section == "overall":
            try:
                key = parts[0]   
                val = float(parts[1])
                overall_metrics[key] = round(val, 4)
            except (ValueError, IndexError):
                continue

        # ── Read per class section 
        elif current_section == "perclass":
            try:
                class_name = parts[0]  
                ap_value   = float(parts[1])
                per_class[class_name] = round(ap_value, 4)
            except (ValueError, IndexError):
                continue

    # ── Validate we got the required metrics 
    required = ["mAP@50", "mAP@50-95", "Precision", "Recall"]
    missing  = [m for m in required if m not in overall_metrics]
    if missing:
        raise HTTPException(
            status_code = 500,
            detail      = f"Missing metrics in results.csv: {missing}"
        )

    return {
        "map50"     : overall_metrics["mAP@50"],
        "map50_95"  : overall_metrics["mAP@50-95"],
        "precision" : overall_metrics["Precision"],
        "recall"    : overall_metrics["Recall"],
        "per_class" : per_class,
        "source"    : "results.csv",
    }


# ─────────────────────────────────────────────────────────────────────────────
#  HELPERS for table insertion
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="D0DCE8"):
    tc        = cell._tc
    tcPr      = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_text(cell, text, bold=False, color=None, font_size=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    para          = cell.paragraphs[0]
    para.alignment = align
    run            = para.add_run(text)
    run.bold       = bold
    run.font.size  = Pt(font_size)
    if color:
        run.font.color.rgb = color

# ─────────────────────────────────────────────────────────────────────────────
#  BUILD DETECTION RESULTS TABLE  (dynamic colours per condition)
# ─────────────────────────────────────────────────────────────────────────────
def build_results_table(doc, discs):
    col_widths = [Mm(30), Mm(38), Mm(50), Mm(35)]
    table      = doc.add_table(rows=1, cols=4)
    table.style = None

    # Header row
    for i, (cell, header) in enumerate(
        zip(table.rows[0].cells,
            ["DISC LEVEL", "CONDITION", "CONFIDENCE SCORE", "CONFIDENCE %"])
    ):
        set_cell_bg(cell, "0D1B2A")
        set_cell_border(cell, "0D1B2A")
        cell.width = col_widths[i]
        set_cell_text(cell, header, bold=True,
                      color=RGBColor(255, 255, 255), font_size=9)

    # Data rows
    for i, disc in enumerate(discs):
        disc_level = disc.get("disc_level",  "N/A")
        condition  = disc.get("condition",   "Not Detected")
        confidence = float(disc.get("confidence", 0))
        conf_score = "N/A" if confidence == 0 else str(round(confidence, 2))
        conf_pct   = "N/A" if confidence == 0 else f"{round(confidence * 100, 1)}%"

        style  = CONDITION_STYLES.get(condition, CONDITION_STYLES["Not Detected"])
        text_c = style["text"]
        cond_bg = style["bg"]
        row_bg  = "E8F4FD" if i % 2 == 1 else "FFFFFF"

        row   = table.add_row()
        cells = row.cells
        for j, w in enumerate(col_widths):
            cells[j].width = w

        # Disc level
        set_cell_bg(cells[0], row_bg);    set_cell_border(cells[0])
        set_cell_text(cells[0], disc_level, color=RGBColor(26, 38, 54), font_size=10)

        # Condition  ← dynamic colour
        set_cell_bg(cells[1], cond_bg);   set_cell_border(cells[1])
        set_cell_text(cells[1], condition, bold=True, color=text_c, font_size=10)

        # Confidence score  ← same colour as condition
        set_cell_bg(cells[2], row_bg);    set_cell_border(cells[2])
        set_cell_text(cells[2], conf_score, color=text_c, font_size=10)

        # Confidence %
        set_cell_bg(cells[3], row_bg);    set_cell_border(cells[3])
        set_cell_text(cells[3], conf_pct, bold=True,
                      color=RGBColor(26, 38, 54), font_size=10)

    return table

# ─────────────────────────────────────────────────────────────────────────────
#  INSERT IMAGE at placeholder paragraph
#  image_data  → bytes  (raw image bytes)
#  placeholder → string to find in the doc paragraphs
# ─────────────────────────────────────────────────────────────────────────────
def insert_image_at_placeholder(doc, placeholder, image_data, width=Inches(2.8)):
    for para in doc.paragraphs:
        if placeholder in para.text:
            # Clear the placeholder text
            para.clear()
            # Insert the image into that paragraph
            run = para.add_run()
            run.add_picture(io.BytesIO(image_data), width=width)
            return True
    return False

# ─────────────────────────────────────────────────────────────────────────────
#  INSERT TABLE at placeholder paragraph
# ─────────────────────────────────────────────────────────────────────────────
def insert_table_at_placeholder(doc, placeholder, discs):
    for para in doc.paragraphs:
        if placeholder in para.text:
            table = build_results_table(doc, discs)
            para._element.addnext(table._tbl)
            para._element.getparent().remove(para._element)
            return True
    return False

#  ENDPOINTS

@app.get("/health", tags=["General"])
def health():
    """Check if server and model are running."""
    return {
        "status"      : "ok",
        "model_loaded": model is not None,
        "timestamp"   : datetime.utcnow().isoformat(),
    }


@app.post("/predict/full", response_model=FullResponse, tags=["Detection"])

async def predict_full(file: UploadFile = File(...)):

    # ── Validate file type 
    if file.content_type not in ["image/jpeg", "image/png", "image/jpg"]:
        raise HTTPException(
            status_code = 400,
            detail      = "Please upload a JPEG or PNG image."
        )

    start = time.time()

    # ── Step 1: uploaded bytes → numpy array 
    file_bytes = await file.read()
    img        = decode_image(file_bytes)
    if img is None:
        raise HTTPException(
            status_code = 400,
            detail      = "Could not read the image. Upload a valid JPEG or PNG."
        )

    # ── Step 2: run YOLOv8 — detection happens here
    detections = run_inference(img)

    # ── Step 3: draw boxes on image using model coordinates 
    vis = draw_boxes(img, detections)

    # ── Step 4: vis (numpy array) → base64 string 
    annotated_b64 = numpy_to_base64(vis)

    # ── Step 5: build text report 
    report = build_report(detections, file.filename, time.time() - start)

    # ── Step 6: return both in one JSON response 
    return {
        "report"         : report,
        "annotated_image": annotated_b64,
    }


@app.get("/metrics", response_model=MetricsResponse, tags=["Evaluation"])
def get_metrics():
    return load_metrics_from_csv(RESULTS_CSV)


@app.post("/generate-report")
async def generate_report(
    data           : str        = Form(...), 
    original_image : UploadFile = File(...), 
):
    uid         = uuid.uuid4().hex[:8].upper()
    filled_docx = TEMP_DIR / f"{uid}_report.docx"
    filled_pdf  = TEMP_DIR / f"{uid}_report.pdf"

    # ── Parse the data object ────────────────────────────────────────────────
    obj     = json.loads(data)
    patient = obj["patient"]
    report  = obj["report"]
    inner   = report["report"]
    discs   = inner["discs"]
    summary = inner["summary"]
    overall = inner["overall_status"]

    # ── Annotated image — decode from base64 ────────────────────────────────
    annotated_b64  = report.get("annotated_image", "")
    if "," in annotated_b64:
        annotated_b64 = annotated_b64.split(",", 1)[1]   # strip  data:image/jpeg;base64,
    annotated_bytes = base64.b64decode(annotated_b64)

    # ── Original image — read uploaded file bytes ────────────────────────────
    original_bytes = await original_image.read()

    # ── Format timestamp ─────────────────────────────────────────────────────
    try:
        dt             = datetime.fromisoformat(inner["timestamp"])
        formatted_date = dt.strftime("%Y-%m-%d %H:%M:%S")
    except:
        formatted_date = inner.get("timestamp", "N/A")

    # ── Summary counts ───────────────────────────────────────────────────────
    h     = summary.get("Herniation",   0)
    b     = summary.get("Bulging",      0)
    n     = summary.get("Normal",       0)
    nd    = summary.get("Not_Detected", 0)
    total = len(discs)

    abnormal = h + b
    status_description = (
        f"The AI model identified abnormalities in {abnormal} of {total} analysed disc level(s): "
        f"{h} disc(s) show signs of Herniation; {b} disc(s) show Bulging. "
        f"Clinical correlation and radiologist review are strongly recommended."
        if abnormal > 0 else
        f"The AI model found no abnormalities across all {total} analysed disc levels. All discs appear Normal."
    )

    # ── STEP 1: Fill text placeholders with docxtpl ──────────────────────────
    context = {
        "detection_id"     : uid,
        "report_generated" : datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "detection_date"   : formatted_date,
        "ai_model"         : "YOLOv8s",

        "patient_name"   : patient.get("name",   "N/A"),
        "patient_dob"    : patient.get("dob",    "N/A"),
        "patient_age"    : str(patient.get("age",    "N/A")),
        "patient_gender" : patient.get("gender", "N/A"),
        "patient_weight" : str(patient.get("weight", "N/A")) + " kg",

        "original_image_filename" : inner.get("image_name", "original.jpg"),
        "detected_image_filename" : inner.get("image_name", "result.jpg"),

        "sum_total"       : str(total),
        "sum_normal"      : str(n),
        "sum_bulging"     : str(b),
        "sum_herniation"  : str(h),
        "sum_notdetected" : str(nd),

        "status_icon"        : overall.upper(),
        "status_title"       : overall,
        "status_description" : status_description,
    }

    tpl = DocxTemplate(TEMPLATE_PATH)
    tpl.render(context)
    tpl.save(str(filled_docx))

    # ── STEP 2: Re-open and inject dynamic table + images ────────────────────
    doc = Document(str(filled_docx))

    # Insert coloured results table
    insert_table_at_placeholder(doc, "results_table_placeholder", discs)

    # Insert original MRI image
    insert_image_at_placeholder(doc, "original_image_placeholder", original_bytes, width=Inches(2.8))

    # Insert annotated result image
    insert_image_at_placeholder(doc, "annotated_image_placeholder", annotated_bytes, width=Inches(2.8))

    doc.save(str(filled_docx))

    # ── STEP 3: Convert .docx → .pdf ────────────────────────────────────────
    convert(str(filled_docx), str(filled_pdf))
    filled_docx.unlink(missing_ok=True)

    return FileResponse(
        path       = str(filled_pdf),
        media_type = "application/pdf",
        filename   = f"spine_report_{patient.get('name', 'report').replace(' ', '_')}.pdf",
    )

#  RUN SERVER
if __name__ == "__main__":
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)